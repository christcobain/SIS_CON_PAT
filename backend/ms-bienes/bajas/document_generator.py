import io
import os
from docx2pdf import convert
from pathlib import Path
from datetime import date
from django.conf import settings
from django.utils import timezone

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


MEDIA_BAJAS      = Path(settings.MEDIA_ROOT) / 'bajas'
MEDIA_BAJAS_PDF  = MEDIA_BAJAS / 'pdfs'
MEDIA_BAJAS_DOCX = MEDIA_BAJAS / 'docx'

LOGO_PATH = Path(settings.MEDIA_ROOT) / 'static' / 'logo_csjln.png'

NEGRO    = RGBColor.from_string('000000')
BLANCO   = RGBColor.from_string('FFFFFF')
GUINDA   = RGBColor.from_string('7F1D1D')
GRIS_OSC = RGBColor.from_string('374151')

GUINDA_HX = '7F1D1D'
GRIS_BORD = 'CCCCCC'
FILA_PAR  = 'F8F8F8'


_SEPARADOR_ANALISIS = '\n§§§\n'

MESES_ES = {
    1: 'enero', 2: 'febrero', 3: 'marzo', 4: 'abril',
    5: 'mayo', 6: 'junio', 7: 'julio', 8: 'agosto',
    9: 'septiembre', 10: 'octubre', 11: 'noviembre', 12: 'diciembre',
}

def _ensure_dirs():
    MEDIA_BAJAS_PDF.mkdir(parents=True, exist_ok=True)
    MEDIA_BAJAS_DOCX.mkdir(parents=True, exist_ok=True)

def _fecha_en_espanol(dt) -> str:
    d = timezone.localtime(dt) if dt else None
    if not d:
        hoy = date.today()
        return f'{hoy.day} de {MESES_ES[hoy.month]} de {hoy.year}'
    return f'{d.day} de {MESES_ES[d.month]} de {d.year}'

def _set_line_spacing_single(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:line'),     '240')
    spacing.set(qn('w:lineRule'), 'auto')
    pPr.append(spacing)

def _set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def _cell_borders(table):
    for row in table.rows:
        for cell in row.cells:
            tc    = cell._tc
            tcPr  = tc.get_or_add_tcPr()
            tcBdr = OxmlElement('w:tcBorders')
            for side in ('top', 'left', 'bottom', 'right'):
                b = OxmlElement(f'w:{side}')
                b.set(qn('w:val'),   'single')
                b.set(qn('w:sz'),    '4')
                b.set(qn('w:space'), '0')
                b.set(qn('w:color'), GRIS_BORD)
                tcBdr.append(b)
            tcPr.append(tcBdr)

def _remove_all_borders(table):
    for row in table.rows:
        for cell in row.cells:
            tc    = cell._tc
            tcPr  = tc.get_or_add_tcPr()
            tcBdr = OxmlElement('w:tcBorders')
            for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
                b = OxmlElement(f'w:{side}')
                b.set(qn('w:val'), 'none')
                tcBdr.append(b)
            tcPr.append(tcBdr)

def _run(paragraph, text: str, *, bold=False, size=10, color=None, italic=False):
    run            = paragraph.add_run(str(text))
    run.bold       = bold
    run.italic     = italic
    run.font.size  = Pt(size)
    run.font.color.rgb = color if color else NEGRO
    return run

def _paragraph(doc, text: str, *, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
               bold=False, size=10, color=None, space_before=2, space_after=4,
               left_indent=None, italic=False) -> object:
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if left_indent is not None:
        p.paragraph_format.left_indent = Cm(left_indent)
    _set_line_spacing_single(p)
    if text:
        _run(p, text, bold=bold, size=size, color=color, italic=italic)
    return p

def _section_title(doc, numero: str, titulo: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(f'{numero}. {titulo.upper()}')
    run.bold           = True
    run.underline      = False
    run.font.size      = Pt(11)
    run.font.color.rgb = NEGRO
    return p

def _subsection_title(doc, numero: str, titulo: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(f'{numero} {titulo.upper()}')
    run.bold           = True
    run.font.size      = Pt(10)
    run.font.color.rgb = NEGRO
    return p

def _narrative_text(doc, text: str):
    if not text or not text.strip():
        _paragraph(doc, '—')
        return
    for linea in text.strip().split('\n'):
        linea = linea.strip()
        if not linea:
            continue
        _paragraph(doc, linea)

def _add_header_line(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    pPr2 = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    btm  = OxmlElement('w:bottom')
    btm.set(qn('w:val'),   'single')
    btm.set(qn('w:sz'),    '6')
    btm.set(qn('w:space'), '1')
    btm.set(qn('w:color'), '000000')
    pBdr.append(btm)
    pPr2.append(pBdr)
    return p

def _split_analisis(analisis: str):
    if not analisis:
        return '', ''
    partes = analisis.split(_SEPARADOR_ANALISIS, 1)
    obs  = partes[0].strip() if len(partes) > 0 else ''
    sust = partes[1].strip() if len(partes) > 1 else ''
    return obs, sust

def generar_docx_baja(baja) -> str:
    _ensure_dirs()

    doc     = Document()
    section = doc.sections[0]
    section.page_width    = Cm(21)
    section.page_height   = Cm(29.7)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    for sname in ('Normal', 'Body Text', 'Default Paragraph Font'):
        try:
            sty = doc.styles[sname]
            sty.font.name      = 'Arial'
            sty.font.size      = Pt(10)
            sty.font.color.rgb = NEGRO
        except (KeyError, AttributeError):
            pass

    sede_nombre   = baja.sede_elabora_nombre or ''
    modulo_nombre = baja.modulo_elabora_nombre or ''

    ANIO_LEMA = '"Año de la Esperanza y el Fortalecimiento de la Democracia"'

    if LOGO_PATH.exists():
        try:
            p_logo = doc.add_paragraph()
            p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_logo = p_logo.add_run()
            run_logo.add_picture(str(LOGO_PATH), width=Inches(1.2))
            p_logo.paragraph_format.space_after = Pt(2)
        except Exception:
            pass

    p_lema = doc.add_paragraph()
    p_lema.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_lema = p_lema.add_run(ANIO_LEMA)
    r_lema.font.size      = Pt(9)
    r_lema.italic         = True
    r_lema.font.color.rgb = NEGRO
    p_lema.paragraph_format.space_after = Pt(2)

    p_header = doc.add_paragraph()
    p_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p_header.add_run('Corte Superior de Justicia de Lima Norte\n')
    r1.bold           = True
    r1.font.size      = Pt(12)
    r1.font.color.rgb = NEGRO
    texto_cabecera = f"Módulo {baja.modulo_elabora_nombre}\nSede {baja.sede_elabora_nombre}"
    run_h = p_header.add_run(texto_cabecera)
    run_h.font.size = Pt(9)
    run_h.font.name = 'Arial'
    p_header.paragraph_format.space_after = Pt(8)

    fecha_str = _fecha_en_espanol(baja.fecha_registro)
    p_fecha           = doc.add_paragraph()
    p_fecha.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rf = p_fecha.add_run(f'Lima, {fecha_str}')
    rf.font.size      = Pt(10)
    rf.font.color.rgb = NEGRO
    p_fecha.paragraph_format.space_after = Pt(6)

    p_tit = doc.add_paragraph()
    r_t   = p_tit.add_run(f'INFORME N° {baja.numero_informe}')
    r_t.bold           = True
    r_t.font.size      = Pt(11)
    r_t.font.color.rgb = NEGRO
    p_tit.paragraph_format.space_after = Pt(6)

    cargo_elabora     = baja.cargo_elabora or 'Asistente de Informática'
    partes            = [p for p in [cargo_elabora, modulo_nombre, sede_nombre] if p]
    cargo_sede_modulo = ' — '.join(partes) if partes else 'Asistente de Informática'

    encab = doc.add_table(rows=4, cols=2)
    encab.alignment = WD_TABLE_ALIGNMENT.LEFT
    etiquetas = ['A', 'De', 'Asunto', 'Referencia']
    mnt_ref = ''
    for det in baja.detalles.all():
        if det.mantenimiento_id and hasattr(det, 'mantenimiento') and det.mantenimiento:
            mnt_ref = det.mantenimiento.numero_orden
            break

    ref_val = baja.numero_informe
    if mnt_ref:
        ref_val = f'{baja.numero_informe} — MNT: {mnt_ref}'

    valores = [
        f'{baja.nombre_destino or "—"}\n{baja.cargo_destino or ""}',
        f'{baja.nombre_elabora or "—"}\n{cargo_sede_modulo}',
        f'Informe técnico de baja de bienes patrimoniales — {sede_nombre}',
        ref_val,
    ]
    col_w = [Cm(3.0), Cm(13.5)]
    for i, (etiq, val) in enumerate(zip(etiquetas, valores)):
        row = encab.rows[i]
        row.cells[0].width = col_w[0]
        row.cells[1].width = col_w[1]
        p0 = row.cells[0].paragraphs[0]
        r0 = p0.add_run(etiq)
        r0.bold            = True
        r0.font.size       = Pt(10)
        r0.font.color.rgb  = NEGRO
        p1 = row.cells[1].paragraphs[0]
        r1b = p1.add_run(': ')
        r1b.bold           = True
        r1b.font.size      = Pt(10)
        r1b.font.color.rgb = NEGRO
        rv = p1.add_run(val)
        rv.font.size       = Pt(10)
        rv.font.color.rgb  = NEGRO
        _set_line_spacing_single(p0)
        _set_line_spacing_single(p1)
    _remove_all_borders(encab)

    _add_header_line(doc)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    p_intro = doc.add_paragraph()
    p_intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    ri = p_intro.add_run(
        'Tengo el agrado de dirigirme a usted, en atención al asunto y en relación '
        'al documento de la referencia, para informarle lo siguiente:'
    )
    ri.font.size      = Pt(10)
    ri.font.color.rgb = NEGRO
    _set_line_spacing_single(p_intro)
    p_intro.paragraph_format.space_after = Pt(8)

    _section_title(doc, '1', 'Antecedentes')
    _narrative_text(doc, baja.antecedentes)

    _section_title(doc, '2', 'Análisis')
    detalles = list(baja.detalles.select_related('motivo_baja', 'mantenimiento').all())

    if detalles:
        p_pre = doc.add_paragraph()
        p_pre.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        rp = p_pre.add_run('Se procedió con la inspección técnica de los siguientes bienes:')
        rp.font.size      = Pt(10)
        rp.font.color.rgb = NEGRO
        _set_line_spacing_single(p_pre)
        p_pre.paragraph_format.space_after = Pt(4)

        headers = ['Tipo', 'Marca', 'Modelo', 'N° Serie', 'Cód. Patrimonial', 'Estado', 'Motivo Baja']
        col_ws  = [Cm(3.2), Cm(2.1), Cm(2.1), Cm(2.1), Cm(2.8), Cm(2.4), Cm(2.6)]
        tb = doc.add_table(rows=1, cols=len(headers))
        tb.alignment = WD_TABLE_ALIGNMENT.LEFT

        hrow = tb.rows[0]
        for j, (hdr, w) in enumerate(zip(headers, col_ws)):
            cell = hrow.cells[j]
            cell.width = w
            _set_cell_bg(cell, GUINDA_HX)
            p_h = cell.paragraphs[0]
            p_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            rh = p_h.add_run(hdr)
            rh.bold           = True
            rh.font.size      = Pt(8)
            rh.font.color.rgb = BLANCO

        for idx, det in enumerate(detalles):
            fila = tb.add_row()
            bg   = FILA_PAR if idx % 2 == 0 else 'FFFFFF'
            vals = [
                det.tipo_bien_nombre,
                det.marca_nombre,
                det.modelo,
                det.numero_serie,
                det.codigo_patrimonial,
                det.estado_funcionamiento,
                getattr(det.motivo_baja, 'nombre', '—'),
            ]
            for j, (val, w) in enumerate(zip(vals, col_ws)):
                cell = fila.cells[j]
                cell.width = w
                _set_cell_bg(cell, bg)
                p_c = cell.paragraphs[0]
                p_c.alignment = WD_ALIGN_PARAGRAPH.CENTER
                rc2 = p_c.add_run(val or '—')
                rc2.font.size      = Pt(8)
                rc2.font.color.rgb = NEGRO

        _cell_borders(tb)
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    _subsection_title(doc, '2.1.', 'Observaciones')

    for det in detalles:
        tiene_obs = any([det.diagnostico_inicial, det.trabajo_realizado, det.estado_funcionamiento])
        if tiene_obs:
            p_bid = doc.add_paragraph()
            p_bid.paragraph_format.space_before = Pt(6)
            p_bid.paragraph_format.space_after  = Pt(2)
            rb = p_bid.add_run(f'Bien: {det.tipo_bien_nombre} — {det.codigo_patrimonial}')
            rb.bold           = True
            rb.font.size      = Pt(10)
            rb.font.color.rgb = NEGRO
            _set_line_spacing_single(p_bid)

            if det.estado_funcionamiento:
                p_ef = doc.add_paragraph()
                p_ef.paragraph_format.left_indent = Cm(0.5)
                p_ef.paragraph_format.space_after = Pt(2)
                re1 = p_ef.add_run('Estado de funcionamiento: ')
                re1.bold = True; re1.font.size = Pt(10); re1.font.color.rgb = NEGRO
                re2 = p_ef.add_run(det.estado_funcionamiento)
                re2.font.size = Pt(10); re2.font.color.rgb = NEGRO
                _set_line_spacing_single(p_ef)

            if det.diagnostico_inicial:
                p_di = doc.add_paragraph()
                p_di.paragraph_format.left_indent = Cm(0.5)
                p_di.paragraph_format.space_after = Pt(2)
                rd1 = p_di.add_run('Diagnóstico inicial: ')
                rd1.bold = True; rd1.font.size = Pt(10); rd1.font.color.rgb = NEGRO
                rd2 = p_di.add_run(det.diagnostico_inicial)
                rd2.font.size = Pt(10); rd2.font.color.rgb = NEGRO
                _set_line_spacing_single(p_di)

            if det.trabajo_realizado:
                p_tr = doc.add_paragraph()
                p_tr.paragraph_format.left_indent = Cm(0.5)
                p_tr.paragraph_format.space_after = Pt(2)
                rt1 = p_tr.add_run('Trabajo realizado: ')
                rt1.bold = True; rt1.font.size = Pt(10); rt1.font.color.rgb = NEGRO
                rt2 = p_tr.add_run(det.trabajo_realizado)
                rt2.font.size = Pt(10); rt2.font.color.rgb = NEGRO
                _set_line_spacing_single(p_tr)

    obs_texto, sust_texto = _split_analisis(baja.analisis)
    if obs_texto:
        for linea in obs_texto.split('\n'):
            linea = linea.strip()
            if not linea:
                continue
            p_an = doc.add_paragraph()
            p_an.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_an.paragraph_format.left_indent = Cm(0.5)
            p_an.paragraph_format.space_after = Pt(2)
            ran = p_an.add_run(linea)
            ran.font.size = Pt(10); ran.font.color.rgb = NEGRO
            _set_line_spacing_single(p_an)

    _subsection_title(doc, '2.2.', 'Sustento Técnico')

    for det in detalles:
        tiene_sust = any([det.diagnostico_final, det.observacion_tecnica])
        if tiene_sust:
            p_bid2 = doc.add_paragraph()
            p_bid2.paragraph_format.space_before = Pt(6)
            p_bid2.paragraph_format.space_after  = Pt(2)
            rb2 = p_bid2.add_run(f'Bien: {det.tipo_bien_nombre} — {det.codigo_patrimonial}')
            rb2.bold = True; rb2.font.size = Pt(10); rb2.font.color.rgb = NEGRO
            _set_line_spacing_single(p_bid2)

            if det.diagnostico_final:
                p_df = doc.add_paragraph()
                p_df.paragraph_format.left_indent = Cm(0.5)
                p_df.paragraph_format.space_after = Pt(2)
                rdf1 = p_df.add_run('Diagnóstico final: ')
                rdf1.bold = True; rdf1.font.size = Pt(10); rdf1.font.color.rgb = NEGRO
                rdf2 = p_df.add_run(det.diagnostico_final)
                rdf2.font.size = Pt(10); rdf2.font.color.rgb = NEGRO
                _set_line_spacing_single(p_df)

            if det.estado_funcionamiento:
                p_eff = doc.add_paragraph()
                p_eff.paragraph_format.left_indent = Cm(0.5)
                p_eff.paragraph_format.space_after = Pt(2)
                reff1 = p_eff.add_run('Estado de funcionamiento final: ')
                reff1.bold = True; reff1.font.size = Pt(10); reff1.font.color.rgb = NEGRO
                reff2 = p_eff.add_run(det.estado_funcionamiento)
                reff2.font.size = Pt(10); reff2.font.color.rgb = NEGRO
                _set_line_spacing_single(p_eff)

            if det.observacion_tecnica:
                p_ot = doc.add_paragraph()
                p_ot.paragraph_format.left_indent = Cm(0.5)
                p_ot.paragraph_format.space_after = Pt(2)
                rot1 = p_ot.add_run('Observación técnica: ')
                rot1.bold = True; rot1.font.size = Pt(10); rot1.font.color.rgb = NEGRO
                rot2 = p_ot.add_run(det.observacion_tecnica)
                rot2.font.size = Pt(10); rot2.font.color.rgb = NEGRO
                _set_line_spacing_single(p_ot)

    if sust_texto:
        for linea in sust_texto.split('\n'):
            linea = linea.strip()
            if not linea:
                continue
            p_st = doc.add_paragraph()
            p_st.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_st.paragraph_format.left_indent = Cm(0.5)
            p_st.paragraph_format.space_after = Pt(2)
            rst = p_st.add_run(linea)
            rst.font.size = Pt(10); rst.font.color.rgb = NEGRO
            _set_line_spacing_single(p_st)

    _section_title(doc, '3', 'Conclusiones')
    _narrative_text(doc, baja.conclusiones)

    _section_title(doc, '4', 'Recomendaciones')
    _narrative_text(doc, baja.recomendaciones)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    p_cierre = doc.add_paragraph()
    p_cierre.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    rc2 = p_cierre.add_run('Es todo cuanto informo a usted para los fines pertinentes.')
    rc2.font.size = Pt(10); rc2.font.color.rgb = NEGRO
    _set_line_spacing_single(p_cierre)
    p_cierre.paragraph_format.space_after = Pt(4)

    p_at = doc.add_paragraph()
    rat = p_at.add_run('Atentamente.')
    rat.font.size = Pt(10); rat.font.color.rgb = NEGRO
    _set_line_spacing_single(p_at)
    p_at.paragraph_format.space_after = Pt(30)

    tf = doc.add_table(rows=3, cols=1)
    tf.alignment = WD_TABLE_ALIGNMENT.CENTER
    _remove_all_borders(tf)

    top_cell_tc = tf.rows[0].cells[0]._tc.get_or_add_tcPr()
    tcBdr2  = OxmlElement('w:tcBorders')
    top_el  = OxmlElement('w:top')
    top_el.set(qn('w:val'),   'single')
    top_el.set(qn('w:sz'),    '4')
    top_el.set(qn('w:color'), '000000')
    tcBdr2.append(top_el)
    top_cell_tc.append(tcBdr2)

    p_nom = tf.rows[0].cells[0].paragraphs[0]
    p_nom.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rn = p_nom.add_run(
        (baja.nombre_elabora or '').upper() or '___________________________'
    )
    rn.bold = True; rn.font.size = Pt(10); rn.font.color.rgb = NEGRO
    _set_line_spacing_single(p_nom)
    p_nom.paragraph_format.space_before = Pt(2)
    p_nom.paragraph_format.space_after  = Pt(1)

    p_crg = tf.rows[1].cells[0].paragraphs[0]
    p_crg.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rc3 = p_crg.add_run(baja.cargo_elabora or 'Asistente de Informática')
    rc3.font.size = Pt(10); rc3.font.color.rgb = NEGRO
    _set_line_spacing_single(p_crg)
    p_crg.paragraph_format.space_before = Pt(1)
    p_crg.paragraph_format.space_after  = Pt(1)

    p_sede_f = tf.rows[2].cells[0].paragraphs[0]
    p_sede_f.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = p_sede_f.add_run(baja.sede_elabora_nombre or '')
    rs.font.size = Pt(10); rs.font.color.rgb = NEGRO
    _set_line_spacing_single(p_sede_f)
    p_sede_f.paragraph_format.space_before = Pt(1)
    p_sede_f.paragraph_format.space_after  = Pt(2)

    ts = (
        timezone.localtime(baja.fecha_registro).strftime('%Y%m%d%H%M%S')
        if baja.fecha_registro
        else timezone.now().strftime('%Y%m%d%H%M%S')
    )
    filename = f'BAJ-{baja.pk}-{ts}.docx'
    docx_abs = MEDIA_BAJAS_DOCX / filename
    doc.save(str(docx_abs))
    return str(docx_abs)


def convertir_docx_a_pdf(docx_abs_path: str) -> str:
    pdf_dir  = str(MEDIA_BAJAS_PDF)
    base     = Path(docx_abs_path).stem
    pdf_path = str(Path(pdf_dir) / f'{base}.pdf')
    try:
        convert(docx_abs_path, pdf_path)
        return pdf_path
    except Exception as e:
        raise RuntimeError(f'Error al convertir PDF: {str(e)}')


def generar_documentos_baja(baja) -> dict:
    """
    Genera el DOCX y convierte a PDF.
    Retorna rutas absolutas/relativas Y los bytes de ambos archivos
    para que _regenerar_documentos() los suba a Supabase.
    """
    docx_abs = generar_docx_baja(baja)
    pdf_abs  = convertir_docx_a_pdf(docx_abs)

    # Leer bytes para subir a Supabase Storage
    with open(docx_abs, 'rb') as f:
        docx_bytes = f.read()
    with open(pdf_abs, 'rb') as f:
        pdf_bytes = f.read()

    media_root = str(settings.MEDIA_ROOT)
    docx_rel   = docx_abs.replace(media_root, '').lstrip(os.sep)
    pdf_rel    = pdf_abs.replace(media_root, '').lstrip(os.sep)

    return {
        'docx_path':  docx_rel,
        'pdf_path':   pdf_rel,
        'docx_bytes': docx_bytes,
        'pdf_bytes':  pdf_bytes,
    }